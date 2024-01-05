from openai import OpenAI
from docx import Document
API_KEY=  "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"
client = OpenAI(api_key=API_KEY)

def create_assistant():
    """
    Create an assistant with retrieval tool activated.
    
    Returns:
        assistant: An assistant object.
        
    >>> isinstance(create_assistant(), openai.Assistant)
    True
    """
    # Create an assistant with retrieval tool activated
    assistant = client.beta.assistants.create(
        instructions="You are a customer support chatbot. Use your knowledge base to best respond to customer queries.",
        model="gpt-3.5-turbo",
        tools=[{"type": "retrieval"}],
    )
    return assistant

def upload_file(file_path):
    """
    Upload a file and add it to the assistant.
    
    Args:
        file_path (str): The path of the file to be uploaded.
        
    Returns:
        file: A file object.
        
    >>> isinstance(upload_file("knowledge.pdf"), openai.File)
    True
    """
    # Upload a file and add it to the assistant
    with open(file_path, "rb") as f:
        file = client.files.create(
            file=f,
            purpose='assistants'
        )
    return file

def use_assistant(assistant, file, query):
    """
    Use the assistant in a conversation, passing in a new file.
    
    Args:
        assistant (openai.Assistant): The assistant to be used.
        file (openai.File): The file to be passed in.
        query (str): The query to be asked.
        
    Returns:
        message: A message object.
        
    >>> isinstance(use_assistant(create_assistant(), upload_file("knowledge.pdf"), "How to turn off this device?"), openai.Message)
    True
    """
    # Use the assistant in a conversation, passing in a new file
    message = client.beta.threads.messages.create(
        thread_id=assistant.id,
        role="user",
        content=query,
        file_ids=[file.id]
    )
    return message

def output_to_docx(message, file_path):
    """
    Output the assistant's response to a Word document.
    
    Args:
        message (openai.Message): The message object to be output.
        file_path (str): The path of the Word document to be output.
        
    Returns:
        None
    """
    # Output the assistant's response to a Word document
    document = Document()
    document.add_heading('Response')
    document.add_paragraph(message.choices[0].text)
    document.save(file_path)

if __name__ == "__main__":
    client = OpenAI(api_key=API_KEY)
    assistant = create_assistant()
    file = upload_file("knowledge.pdf")
    message = use_assistant(assistant, file, "I can not find in the PDF manual how to turn off this device.")
    output_to_docx(message, "output.docx")